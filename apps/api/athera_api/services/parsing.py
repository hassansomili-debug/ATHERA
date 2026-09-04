"""تفكيك المستندات | Document parsing (§33.1).

الهدف ليس استخراج نص، بل إنتاج **مقاطع قابلة للاستشهاد**: لكل مقطع صفحة
وقسم وفقرة، لأن حقيقة بلا موضع لا تصلح مصدرًا (§10.2، §29.2).

التقطيع بنيوي لا بحجم ثابت: نحترم حدود الفقرات والعناوين، ثم نقسّم الفقرة
الطويلة عند حدود الجمل.
"""
from __future__ import annotations

import re
from dataclasses import dataclass

MAX_CHUNK_CHARS = 1_800
MIN_CHUNK_CHARS = 40
MAX_CHUNKS_PER_FILE = 5_000

# عناوين معروفة بالاسم — عربية وإنجليزية.
_HEADING_PATTERNS = [
    re.compile(r"^\s*(\d+(?:\.\d+)*)\s+(\S.{0,120})$"),
    re.compile(
        r"^\s*(الفصل|المبحث|القسم|المقدمة|الخاتمة|المراجع|الملاحق|السيرة|المهارات|"
        r"الخبرات|المؤهلات|التدريب|الأبحاث|المنشورات)\b.{0,120}$"
    ),
    re.compile(
        r"^\s*(abstract|introduction|literature review|methodology|method|results|"
        r"discussion|conclusion|references|appendix|education|experience|publications|"
        r"skills|training|curriculum vitae)\b.{0,120}$",
        re.IGNORECASE,
    ),
]

# علامات نهاية جملة — العنوان لا ينتهي بها.
_TERMINAL_PUNCTUATION = ".!?؟،؛:"

_SENTENCE_SPLIT = re.compile(r"(?<=[.!?؟।۔])\s+|(?<=\.)\s+")


@dataclass(slots=True)
class ParsedChunk:
    seq: int
    text: str
    locator: str
    page_number: int | None
    section_path: str | None
    paragraph_index: int | None

    @property
    def char_count(self) -> int:
        return len(self.text)


class UnsupportedDocument(Exception):
    """نوع لا نستطيع تفكيكه — نقولها صراحةً بدل إنتاج نص فارغ صامت."""


class NoTextLayer(UnsupportedDocument):
    """مستندٌ سليم بلا طبقة نصّ — **وهذه ليست «نوعًا غير مدعوم»**.

    الفرق ليس تصنيفًا: مَن رفع ملفًّا لا يقرؤه المفكِّك يُقال له إن نوعه
    غير مدعوم، ومَن رفع رسالةً ممسوحةً ضوئيًّا مستندُه مدعومٌ تمامًا —
    والذي ينقصه أن يُقرأ ضوئيًّا. وطيُّ الحالين في رسالةٍ واحدة يجعل
    الشاشة تعرض «أعد المحاولة» على مسحٍ ضوئيّ، فتَعِد بما لا يقع: إعادة
    القراءة تُنتج النتيجة نفسها حرفًا بحرف.

    وهي فرعٌ من `UnsupportedDocument` عمدًا: كلُّ `except` قائم في المستودع
    يبقى يمسكها كما كان، والذي يُضاف هو **تمييزُ** من أراد التمييز.
    """


def _is_heading(line: str, *, is_whole_paragraph: bool = False) -> bool:
    """عنوان بالاسم، أو بالبنية.

    الاعتماد على قائمة كلمات وحدها يسقط أمام أي مستند حقيقي: كل جامعة تكتب
    عناوينها بصيغتها. لذا نضيف قاعدة بنيوية — سطر قصير، وحده في فقرته، بلا
    علامة نهاية جملة — وهي ما يلتقط «السيرة الذاتية» و«المهارات» وأمثالها.
    """
    stripped = line.strip()
    if not stripped or len(stripped) > 140:
        return False
    if any(pattern.match(stripped) for pattern in _HEADING_PATTERNS):
        return True
    return (
        is_whole_paragraph
        and len(stripped) <= 60
        and len(stripped.split()) <= 8
        and stripped[-1] not in _TERMINAL_PUNCTUATION
    )


def _split_long(text: str) -> list[str]:
    """يقسّم فقرة طويلة عند حدود الجمل، لا في منتصف كلمة."""
    if len(text) <= MAX_CHUNK_CHARS:
        return [text]
    parts, current = [], ""
    for sentence in _SENTENCE_SPLIT.split(text):
        if not sentence:
            continue
        if len(current) + len(sentence) + 1 > MAX_CHUNK_CHARS and current:
            parts.append(current.strip())
            current = sentence
        else:
            current = f"{current} {sentence}".strip()
    if current.strip():
        parts.append(current.strip())
    return parts


def _build(pages: list[tuple[int | None, str]]) -> list[ParsedChunk]:
    chunks: list[ParsedChunk] = []
    section = None
    seq = 0

    for page_number, page_text in pages:
        paragraph_index = 0
        for raw_paragraph in re.split(r"\n\s*\n", page_text):
            paragraph = re.sub(r"[ \t]+", " ", raw_paragraph).strip()
            if not paragraph:
                continue

            first_line = paragraph.split("\n", 1)[0].strip()
            whole = len(paragraph.strip()) <= len(first_line) + 2
            if _is_heading(first_line, is_whole_paragraph=whole):
                section = first_line
                if len(paragraph) <= len(first_line) + 2:
                    continue  # عنوان وحده ليس مقطعًا

            paragraph_index += 1
            for part in _split_long(paragraph.replace("\n", " ")):
                if len(part) < MIN_CHUNK_CHARS:
                    continue
                seq += 1
                if seq > MAX_CHUNKS_PER_FILE:
                    return chunks
                locator_bits = []
                if page_number is not None:
                    locator_bits.append(f"p.{page_number}")
                if section:
                    locator_bits.append(f"§{section[:60]}")
                locator_bits.append(f"¶{paragraph_index}")
                chunks.append(
                    ParsedChunk(
                        seq=seq,
                        text=part,
                        locator=" ".join(locator_bits),
                        page_number=page_number,
                        section_path=section,
                        paragraph_index=paragraph_index,
                    )
                )
    return chunks


def parse_pdf(data: bytes) -> list[ParsedChunk]:
    try:
        from pypdf import PdfReader  # noqa: PLC0415
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedDocument("pypdf is required to parse PDF files") from exc

    import io

    reader = PdfReader(io.BytesIO(data))
    pages = [(index + 1, page.extract_text() or "") for index, page in enumerate(reader.pages)]
    if not any(text.strip() for _, text in pages):
        # PDF ممسوح ضوئيًا بلا طبقة نص — نقولها بدل إنتاج صفر حقائق صامتًا.
        #
        # **وبصنفٍ مميَّز لا برسالةٍ تُقرأ بـ`grep`.** خطُّ الأنابيب يحتاج أن
        # يفرّق هذه عن «نوعٍ غير مدعوم» ليكتب الحال الصحيحة على الرسالة؛
        # ومطابقةُ نصّ الاستثناء تنكسر بأول تحسينٍ للصياغة.
        raise NoTextLayer(
            "no extractable text layer (scanned PDF); OCR is not available yet"
        )
    return _build(pages)


def parse_docx(data: bytes) -> list[ParsedChunk]:
    try:
        import docx  # noqa: PLC0415
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedDocument("python-docx is required to parse DOCX files") from exc

    import io

    document = docx.Document(io.BytesIO(data))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            blocks.append("")
            continue
        style = (paragraph.style.name or "").lower()
        blocks.append(text if "heading" not in style else f"{text}\n")
    # DOCX بلا ترقيم صفحات موثوق — نترك page_number فارغًا بدل تخمينه.
    return _build([(None, "\n\n".join(blocks))])


def parse_text(data: bytes) -> list[ParsedChunk]:
    return _build([(None, data.decode("utf-8", errors="replace"))])


def can_parse(content_type: str, filename: str = "") -> bool:
    """أيستطيع المفكِّك قراءة هذا النوع؟ — **يُسأل قبل الوعد لا بعده**.

    والمكتبة تعرض «معالجة المستند» لما يمكن معالجته وحده؛ فزرٌّ يُعرض على
    ملفٍ لا يُقرأ يَعِد ثم يخذل. والجواب من `parse` نفسه لا من قائمةٍ ثانية
    تفترق عنه بأول نوعٍ يُضاف.
    """
    lowered = (filename or "").lower()
    return (
        content_type == "application/pdf" or lowered.endswith(".pdf")
        or "wordprocessingml" in content_type or lowered.endswith(".docx")
        or content_type.startswith("text/") or lowered.endswith((".txt", ".md"))
    )


def parse(data: bytes, content_type: str, filename: str = "") -> list[ParsedChunk]:
    lowered = filename.lower()
    if content_type == "application/pdf" or lowered.endswith(".pdf"):
        return parse_pdf(data)
    if "wordprocessingml" in content_type or lowered.endswith(".docx"):
        return parse_docx(data)
    if content_type.startswith("text/") or lowered.endswith((".txt", ".md")):
        return parse_text(data)
    raise UnsupportedDocument(f"unsupported content type for parsing: {content_type}")
